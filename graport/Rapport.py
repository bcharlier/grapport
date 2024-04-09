import datetime
import os.path

import pandas as pd
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import graport
from graport.utils import calculate_age, short_path



class Rapport:
    """
    Classe Rapport.  Un rapport est un document qui contient des informations sur un candidat.
    Lit les données d'un candidat et les intègre dans un document Word.

    Attributs:
    candidatures: Candidatures instance
    rapporteur: Rapporteur instance

    Méthodes:
    generate: Génère un rapport pour chaque poste.
    """

    def __init__(self, candidatures, rapporteur):
        self.candidatures = candidatures
        self.rapporteur = rapporteur

        # verifier si le fichier logo_um.png existe
        if not os.path.exists(os.path.join(graport.config.image_dir, 'logo.png')):
            self.logo = None
        else:
            self.logo = os.path.join(graport.config.image_dir, 'logo.png')

    def generate(self):

        for id in self.rapporteur.rapports:

            candidat = self.candidatures.get_candidate(id, type=self.rapporteur.type)
            print(f"Generating report for candidate {candidat['Nom']} {candidat['Prénom']}...", end="")
            filename = os.path.join(graport.config.data_dir, f'rapport_{candidat["Nom"]}-{candidat["Prénom"]}.docx')
            doc = self.create_document(candidat)
            print('done. Saving... ', end='')
            doc.save(filename)
            print(f"Saved to {short_path(filename)}.")

    def create_document(self, candidat):
        document = Document()

        font = document.styles['Normal'].font
        font.name = 'Helvetica'
        font.size = Pt(10)

        paragraph_format = document.styles['Normal'].paragraph_format
        paragraph_format.space_before = Pt(12)

        if self.logo is not None:
            document.add_picture(os.path.join(graport.config.data_dir, 'logo.png'), width=Inches(1.25))


        title = document.add_heading(level=1)
        title.alignment = 1
        run = title.add_run(f"RAPPORT SUR UNE CANDIDATURE\nN° Galaxie du Poste : {candidat['Référence GALAXIE']}")
        #run.bold = True
        #run.font.size = Inches(0.25)
        #run.font.color.rgb = RGBColor(0, 0, 0)



        document.add_paragraph()

        #---------------------------------------------
        # Rapporteur
        #---------------------------------------------

        reviewer = document.add_paragraph()
        reviewer.add_run('Nom de la rapporteuse : ').bold = True
        reviewer.add_run(f"{self.rapporteur.nom}")
        reviewer.add_run('\n')
        reviewer.add_run('Corps : ').bold = True
        reviewer.add_run(candidat['Corps'])
        reviewer.add_run(' ' * 5)
        reviewer.add_run('Section : ').bold = True
        if candidat['Autre section'] == "":
            reviewer.add_run(f"{candidat['Section']}")
        else:
            reviewer.add_run("/".join([candidat['Section'], candidat['Autre section']]))

        reviewer.add_run(" " * 5)
        reviewer.add_run('Profil : ').bold = True
        reviewer.add_run(candidat['Profil J.o'])

        #---------------------------------------------
        # Candidat
        #---------------------------------------------

        h2 = document.add_heading(level=2)
        run = h2.add_run(f"Candidat N° {candidat['N° candidat']}")

        cand = document.add_paragraph()
        cand.add_run('Nom : ').bold = True
        cand.add_run(candidat['Nom'])
        cand.add_run(' ' * 5)
        cand.add_run("Nom d'usage : ").bold = True
        cand.add_run(candidat["Nom d'usage ou marital"])
        cand.add_run(' ' * 5)
        cand.add_run('Prénom : ').bold = True
        cand.add_run(candidat['Prénom'])

        cand.add_run('\n')
        cand.add_run('Né(e) le : ').bold = True
        cand.add_run(candidat['Né(e) le'] + '(' + calculate_age(candidat['Né(e) le']) + ' ans)')
        cand.add_run(' ' * 5)
        cand.add_run('à : ').bold = True
        cand.add_run(candidat['Lieu de naissance'])

        cand_s = document.add_paragraph()
        cand_s.add_run('Situation professionnelle : ').bold = True
        cand_s.add_run(candidat['Situation professionnelle'])
        cand_s.add_run(' ' * 5)
        cand_s.add_run("Lieu d'exercice : ").bold = True
        cand_s.add_run(candidat["Lieu d'exercice"])

        cand_s.add_run('\n')
        cand_s.add_run('Qualification : ').bold = True
        cand_s.add_run(f"{candidat['N° de qualif']}")

        #---------------------------------------------
        # Diplôme
        #---------------------------------------------

        document.add_heading('Diplôme', level=2)

        titre = document.add_paragraph()
        titre.add_run('Titre : ').bold = True
        titre.add_run(candidat['Titres'])

        these = document.add_paragraph()
        these.add_run('Sujet de Thèse : ').bold = True
        these.add_run(candidat['Titre thèse'])

        these_d = document.add_paragraph()
        these_d.add_run('Date de soutenance : ').bold = True
        these_d.add_run(candidat['Date soutenance'])
        these_d.add_run(' ' * 5)
        these_d.add_run('Etablissement : ').bold = True
        these_d.add_run(candidat['Lieu soutenance'])

        these_j = document.add_paragraph()
        these_j.add_run('Directeur de recherche : ').bold = True
        these_j.add_run(candidat['Directeur Thèse'])
        these_j.add_run('\n')
        these_j.add_run('Jury : ').bold = True
        these_j.add_run(candidat['Jury'])

        if candidat['Autres diplômes'] != "":
            autre = document.add_paragraph()
            autre.add_run('Autres diplômes : ').bold = True
            autre.add_run(f"{candidat['Autres diplômes']}")

        #---------------------------------------------
        # Activités
        #---------------------------------------------

        document.add_heading("Activités", level=2)

        enseignement = document.add_paragraph()
        enseignement.add_run('Activités enseignement : ').bold = True
        enseignement.add_run(candidat['Activités enseignement'])

        recherche = document.add_paragraph()
        recherche.add_run('Activités recherche : ').bold = True
        recherche.add_run(candidat['Activités recherche'])

        administration = document.add_paragraph()
        administration.add_run('Activités administratives : ').bold = True
        administration.add_run(candidat['Activités administratives'])

        document.add_heading("Publications", level=2)

        if candidat['Travaux'] != "":
            publications = document.add_paragraph()
            publications.add_run('Travaux : ').bold = True
            publications.add_run(candidat['Travaux'])


        #---------------------------------------------
        # Proposition et avis
        #---------------------------------------------

        title = document.add_heading(level=1)
        title.alignment = 1
        run = title.add_run("Proposition et avis détaillé de la rapporteuse")

        table = document.add_table(rows=4, cols=1)

        cells = table.rows[0].cells
        par = cells[0].add_paragraph()
        par.alignment = 1
        par.add_run('Compte tenu du profil du poste, tel qu’il a été diffusé, je recommande l’audition du candidat ci-dessus (favorable ou réservé ou défavorable):')

        cells = table.rows[1].cells
        par = cells[0].add_paragraph()
        par.add_run('Avis :').bold = True
        par.add_run('\n')

        cells = table.rows[2].cells
        cells[0].text = 'AVIS DETAILLE DE LA DÉCISION (synthèse et adéquation au profil)'

        cells = table.rows[3].cells

        par = cells[0].add_paragraph()
        par.add_run('Motif : ').bold = True
        par.add_run('\n')

        #---------------------------------------------
        # Signature
        #---------------------------------------------

        document.add_paragraph()
        sig = document.add_paragraph()
        sig.add_run(f'Fait à {self.rapporteur.ville if self.rapporteur.ville is not None else " " * 5}, le {datetime.datetime.now().strftime("%d/%m/%Y")}')
        sig.add_run(' ' * 25)
        sig.add_run(f"Signature {self.rapporteur.nom}")
        if self.rapporteur.signature is not None:
            document.add_picture(self.rapporteur.signature, width=Inches(1))

        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.add_page_break()

        return document

