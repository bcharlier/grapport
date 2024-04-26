import os, datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from grapport.utils import calculate_age

"""
This module contains the template for the document. It creates a Word document with the information of a candidate.
It should contains the following functions:

- create_document: creates the document
- save: saves the document
- extension: extension of the file
- logo: logo of the institution
"""

extension = "docx"
logo = "logo.png" # None si pas de logo

def create_document(candidat, rapporteur):
    document = Document()

    font = document.styles['Normal'].font
    font.name = 'Helvetica'
    font.size = Pt(10)

    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(12)

    if logo is not None:
        document.add_picture(os.path.join(os.path.dirname(__file__), logo), width=Inches(1.25))

    title = document.add_heading(level=1)
    title.alignment = 1
    run = title.add_run(f"RAPPORT SUR UNE CANDIDATURE\nN° Galaxie du Poste : {candidat['Référence GALAXIE']}")
    # run.bold = True
    # run.font.size = Inches(0.25)
    # run.font.color.rgb = RGBColor(0, 0, 0)

    document.add_paragraph()

    # ---------------------------------------------
    # Rapporteur
    # ---------------------------------------------

    reviewer = document.add_paragraph()
    reviewer.add_run('Nom de la rapporteuse : ').bold = True
    reviewer.add_run(f"{rapporteur.nom}")
    reviewer.add_run('\n')
    reviewer.add_run('Corps : ').bold = True
    reviewer.add_run(candidat['Corps'])
    reviewer.add_run(' ' * 5)
    reviewer.add_run('Section : ').bold = True
    if candidat['Autre section'] == "":
        reviewer.add_run(f"{candidat['Section']}")
    else:
        reviewer.add_run("/".join([candidat['Section'], candidat['Autre section']]))

    reviewer.add_run(" " * 5)
    reviewer.add_run('Profil : ').bold = True
    reviewer.add_run(candidat['Profil J.o'])

    # ---------------------------------------------
    # Candidat
    # ---------------------------------------------

    h2 = document.add_heading(level=2)
    run = h2.add_run(f"Candidat N° {candidat['N° candidat']}")

    cand = document.add_paragraph()
    cand.add_run('Nom : ').bold = True
    cand.add_run(candidat['Nom'])
    cand.add_run(' ' * 5)
    cand.add_run("Nom d'usage : ").bold = True
    cand.add_run(candidat["Nom d'usage ou marital"])
    cand.add_run(' ' * 5)
    cand.add_run('Prénom : ').bold = True
    cand.add_run(candidat['Prénom'])

    cand.add_run('\n')
    cand.add_run('Né(e) le : ').bold = True
    cand.add_run(candidat['Né(e) le'] + ' (' + calculate_age(candidat['Né(e) le']) + ' ans)')
    cand.add_run(' ' * 5)
    cand.add_run('à : ').bold = True
    cand.add_run(candidat['Lieu de naissance'])

    cand_s = document.add_paragraph()
    cand_s.add_run('Situation professionnelle : ').bold = True
    cand_s.add_run(candidat['Situation professionnelle'])
    cand_s.add_run(' ' * 5)
    cand_s.add_run("Lieu d'exercice : ").bold = True
    cand_s.add_run(candidat["Lieu d'exercice"])

    cand_s.add_run('\n')
    cand_s.add_run('Qualification : ').bold = True
    cand_s.add_run(candidat['Référence qualif'].split('-')[1])
    cand_s.add_run(' ' * 5)
    cand_s.add_run('Section :').bold = True
    cand_s.add_run(candidat['Référence qualif'].split('-')[2])

    # ---------------------------------------------
    # Diplôme
    # ---------------------------------------------

    document.add_heading('Formation et expérience professionnelle', level=2)

    titre = document.add_paragraph()
    titre.add_run('Titre : ').bold = True
    titre.add_run(candidat['Titres'])

    these = document.add_paragraph()
    these.add_run('Sujet de Thèse : ').bold = True
    these.add_run(candidat['Titre thèse'])

    these_d = document.add_paragraph()
    these_d.add_run('Date de soutenance : ').bold = True
    these_d.add_run(candidat['Date soutenance'])
    these_d.add_run(' ' * 5)
    these_d.add_run('Établissement : ').bold = True
    these_d.add_run(candidat['Lieu soutenance'])

    these_j = document.add_paragraph()
    these_j.add_run('Directeur de recherche : ').bold = True
    these_j.add_run(candidat['Directeur Thèse'])
    these_j.add_run('\n')
    these_j.add_run('Jury : ').bold = True
    these_j.add_run(candidat['Jury'])

    autre = document.add_paragraph()
    autre.add_run('Autres diplômes : ').bold = True
    autre.add_run(f"{candidat['Autres diplômes']}" if candidat['Autres diplômes'] else '.' * 20).bold = False

    exp_r = document.add_paragraph()
    exp_r.add_run('Expérience professionnelle en recherche : ').bold = True
    exp_r.add_run("." * 20).bold = False

    exp_e = document.add_paragraph()
    exp_e.add_run('Expérience professionnelle en enseignement : ').bold = True
    exp_e.add_run("." * 20).bold = False

    # ---------------------------------------------
    # Activités Recherche
    # ---------------------------------------------

    document.add_heading("Activités de recherche", level=2)

    recherche = document.add_paragraph()
    recherche.add_run('Activités recherche : ').bold = True
    recherche.add_run(candidat['Activités recherche'])

    publications = document.add_paragraph()
    publications.add_run('Travaux : ').bold = True
    publications.add_run(candidat['Travaux'] if candidat['Travaux'] else '.' * 20)

    publications = document.add_paragraph()
    publications.add_run('Encadrement : ').bold = True
    publications.add_run('.' * 20).bold = False

    code = document.add_paragraph()
    code.add_run('Logiciels : ').bold = True
    code.add_run('.' * 20).bold = False

    # ---------------------------------------------
    # Activités Enseignement
    # ---------------------------------------------

    document.add_heading("Activités d'enseignement", level=2)

    enseignement = document.add_paragraph()
    enseignement.add_run('Activités enseignement : ').bold = True
    enseignement.add_run(candidat['Activités enseignement'])

    administration = document.add_paragraph()
    administration.add_run('Responsabilité administratives (école doctorales, filières, formation pro) : ').bold = True
    administration.add_run('.' * 20).bold = False

    # ---------------------------------------------
    # Activités collectives
    # ---------------------------------------------

    document.add_heading("Responsabilités collectives", level=2)

    administration = document.add_paragraph()
    administration.add_run('Activités administratives : ').bold = True
    administration.add_run(candidat['Activités administratives'])

    editorial = document.add_paragraph()
    editorial.add_run('Activités éditoriales : ').bold = True
    editorial.add_run('.' * 20).bold = False

    evaluation = document.add_paragraph()
    evaluation.add_run('Participation à des instances d’évaluation : ').bold = True
    evaluation.add_run('.' * 20).bold = False

    mediation = document.add_paragraph()
    mediation.add_run('Médiation scientifique : ').bold = True
    mediation.add_run('.' * 20).bold = False

    # ---------------------------------------------
    # Proposition et avis
    # ---------------------------------------------

    title = document.add_heading(level=1)
    title.alignment = 1
    title.add_run("Proposition et avis détaillé de la rapporteuse")

    table = document.add_table(rows=6, cols=1)

    cells = table.rows[0].cells
    cells[0].text = "Adéquation du profil du candidat au profil d'enseignement"
    p = cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.bold = True

    cells = table.rows[1].cells
    par = cells[0].add_paragraph()
    par.add_run('\n')

    cells = table.rows[2].cells
    cells[0].text = "Adéquation du profil du candidat au profil de recherche"
    p = cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.bold = True

    cells = table.rows[3].cells
    par = cells[0].add_paragraph()
    par.add_run('\n')

    cells = table.rows[4].cells
    cells[0].text = 'Avis (favorable ou réservé ou défavorable)'
    p = cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.bold = True

    cells = table.rows[5].cells
    par = cells[0].add_paragraph()
    par.add_run('\n')

    # ---------------------------------------------
    # Signature
    # ---------------------------------------------

    document.add_paragraph()
    sig = document.add_paragraph()
    sig.add_run(
        f'Fait à {rapporteur.ville if rapporteur.ville is not None else " " * 5}, le {datetime.datetime.now().strftime("%d/%m/%Y")}')
    sig.add_run(' ' * 25)
    sig.add_run(f"Signature, {rapporteur.nom}")
    if rapporteur.signature is not None:
        document.add_picture(rapporteur.signature, width=Inches(1.5))

        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_page_break()

    return document

def save(document, filename):
    document.save(filename)